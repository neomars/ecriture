import zipfile
from werkzeug.utils import secure_filename
from bs4 import BeautifulSoup
from flask import Response
import os
import json
import re






import platform
from packaging import version
import json
import urllib.request
from flask import Flask, jsonify, request, send_file, render_template
from project_manager import NovelProject
from ai_client import AIClient

def get_synonyms(word, lang="fr"):
    """Lookup synonyms from the WOLF synonyms table in lexique.db, utilizing lemma fallback."""
    if lang != "fr":
        return []

    db_path = "lexique.db"
    if not os.path.exists(db_path):
        return []

    w_clean = word.lower().strip(".,!?;:\"'()[]{}«»")
    if not w_clean:
        return []

    import sqlite3
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        # 1. Query synonyms for the clean word itself
        cursor.execute(
            "SELECT DISTINCT synonym FROM synonyms WHERE word = ? LIMIT 20",
            (w_clean,)
        )
        syns = [r[0] for r in cursor.fetchall() if r[0]]

        # 2. Get the lemma of the selected word
        cursor.execute("SELECT lemme FROM lexique WHERE ortho = ? LIMIT 1", (w_clean,))
        row = cursor.fetchone()
        if row and row[0]:
            lemma = row[0].lower().strip()
            if lemma != w_clean:
                # Query synonyms of the lemma
                cursor.execute(
                    "SELECT DISTINCT synonym FROM synonyms WHERE word = ? LIMIT 20",
                    (lemma,)
                )
                for r in cursor.fetchall():
                    if r[0] and r[0] not in syns and r[0] != w_clean:
                        syns.append(r[0])

        conn.close()
        return syns[:20]
    except Exception as e:
        print("Error querying synonyms from lexique.db:", e)
        return []

app = Flask(__name__, template_folder='templates')
ai_client = AIClient()

PROJECTS_DIR = "projects"
ACTIVE_CONFIG_FILE = "active_project.txt"

# Ensure projects directory exists
if not os.path.exists(PROJECTS_DIR):
    os.makedirs(PROJECTS_DIR)

def get_active_project_filename():
    """Gets the currently active project filename from config or defaults."""
    if os.path.exists(ACTIVE_CONFIG_FILE):
        with open(ACTIVE_CONFIG_FILE, 'r', encoding='utf-8') as f:
            fn = f.read().strip()
            if fn and os.path.exists(os.path.join(PROJECTS_DIR, fn)):
                return fn

    # Prioritize the Corneille "Le Cid" example as first-time default if it exists
    if os.path.exists(os.path.join(PROJECTS_DIR, "le_cid_corneille.json")):
        set_active_project_filename("le_cid_corneille.json")
        return "le_cid_corneille.json"

    # Fallback: scan projects directory or create a default one
    files = [f for f in os.listdir(PROJECTS_DIR) if f.endswith(".json")]
    if files:
        # Save first as active
        set_active_project_filename(files[0])
        return files[0]

    # If no files exist, create a default one
    default_fn = "my_novel_project.json"
    default_path = os.path.join(PROJECTS_DIR, default_fn)

    # If we have an existing my_novel_project.json in root, move it or create a clean one
    if os.path.exists("my_novel_project.json"):
        import shutil
        shutil.copy("my_novel_project.json", default_path)
    else:
        proj = NovelProject(default_path)
        proj.save()

    set_active_project_filename(default_fn)
    return default_fn

def set_active_project_filename(filename):
    """Saves the active project filename to the config file."""
    with open(ACTIVE_CONFIG_FILE, 'w', encoding='utf-8') as f:
        f.write(filename)

# Instantiate the active project
active_filename = get_active_project_filename()
project = NovelProject(os.path.join(PROJECTS_DIR, active_filename))

def get_scene_context(scene_id):
    """
    Gathers structured lore context about characters and notes associated with the active scene.
    A character is associated if linked via the checklist or via a plot card in this scene.
    A note is associated if its title is mentioned in the scene's title, content, or plot cards.
    """
    if not scene_id or not project or not project.data:
        return ""

    # 1. Collect associated character IDs
    char_ids = set()
    for char in project.data.get("characters", []):
        if scene_id in char.get("linked_scenes", []):
            char_ids.add(char["id"])

    # Plot cards for this scene
    for card in project.data.get("plot", {}).get("cards", []):
        if card.get("scene_id") == scene_id:
            for c_id in card.get("characters", []):
                char_ids.add(c_id)

    char_name_map = {c["id"]: c.get("name", "Inconnu") for c in project.data.get("characters", [])}

    char_lore_blocks = []
    for char in project.data.get("characters", []):
        if char["id"] in char_ids:
            name = char.get("name", "Inconnu")
            role = char.get("role", "")
            aliases = ", ".join(char.get("aliases", []))
            traits = ", ".join(char.get("traits", []))
            appearance = char.get("appearance", "")
            desc = char.get("description", "")
            notes = char.get("notes", "")

            rel_strs = []
            for rel in char.get("relations", []):
                target_id = rel.get("target_id")
                target_name = char_name_map.get(target_id, "Inconnu")
                rel_type = rel.get("type", "")
                rel_desc = rel.get("description", "")

                parts = []
                if rel_type:
                    parts.append(rel_type)
                if rel_desc:
                    parts.append(rel_desc)
                if parts:
                    rel_strs.append(f"- Relation avec {target_name} : {', '.join(parts)}")

            block = []
            block.append(f"Personnage : {name}")
            if role:
                block.append(f"  Rôle : {role}")
            if aliases:
                block.append(f"  Alias/Surnoms : {aliases}")
            if traits:
                block.append(f"  Traits de caractère : {traits}")
            if appearance:
                block.append(f"  Apparence physique : {appearance}")
            if desc:
                block.append(f"  Description : {desc}")
            if notes:
                block.append(f"  Notes : {notes}")
            if rel_strs:
                block.append("  Relations :\n" + "\n".join(rel_strs))

            char_lore_blocks.append("\n".join(block))

    # 2. Collect associated story notes by scanning scene content, title and plot cards
    scene_title = ""
    scene_content = ""

    def find_scene_text(nodes):
        for n in nodes:
            if n.get("id") == scene_id:
                return n.get("title", ""), n.get("content", "")
            if "children" in n:
                t, c = find_scene_text(n["children"])
                if t or c:
                    return t, c
        return "", ""

    scene_title, scene_content = find_scene_text(project.data.get("manuscript", []))
    combined_scene_text = f"{scene_title}\n{scene_content}".lower()

    for card in project.data.get("plot", {}).get("cards", []):
        if card.get("scene_id") == scene_id:
            combined_scene_text += f"\n{card.get('title', '')}\n{card.get('content', '')}".lower()

    note_lore_blocks = []
    for note in project.data.get("story_notes", []):
        note_title = note.get("title", "")
        if note_title and note_title.lower() in combined_scene_text:
            note_type = note.get("type", "")
            note_content = note.get("content", "")

            block = []
            block.append(f"Note : {note_title}")
            if note_type:
                block.append(f"  Type : {note_type}")
            if note_content:
                block.append(f"  Contenu : {note_content}")
            note_lore_blocks.append("\n".join(block))

    context_parts = []
    if char_lore_blocks:
        context_parts.append("=== LORE DES PERSONNAGES ASSOCIÉS À CETTE SCÈNE ===\n" + "\n\n".join(char_lore_blocks))
    if note_lore_blocks:
        context_parts.append("=== LORE DES NOTES ASSOCIÉES À CETTE SCÈNE ===\n" + "\n\n".join(note_lore_blocks))

    if context_parts:
        return "\n\n".join(context_parts)
    return ""

# Ensure the active project data has a default language setting if not present
if "lang" not in project.data["settings"]:
    project.data["settings"]["lang"] = "en"

@app.route('/')
def index():
    """Renders the main single-page web interface."""
    return render_template('index.html')

@app.route('/api/project', methods=['GET'])
def get_project():
    """Returns the currently active project state."""
    global project
    project.recalculate_word_counts()
    return jsonify(project.data)

@app.route('/api/project', methods=['POST'])
def save_project():
    """Updates the active project state from client data and persists it to JSON."""
    global project
    client_data = request.json
    if not client_data:
        return jsonify({"error": "No data provided"}), 400

    project.data = client_data
    project.save()
    return jsonify({"status": "success", "data": project.data})

@app.route('/api/projects', methods=['GET'])
def list_projects():
    """Lists all available projects in the projects/ directory with their titles."""
    projects_list = []
    files = [f for f in os.listdir(PROJECTS_DIR) if f.endswith(".json")]
    for filename in files:
        filepath = os.path.join(PROJECTS_DIR, filename)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                title = data.get("settings", {}).get("title", filename)
                projects_list.append({"filename": filename, "title": title})
        except Exception as e:
            # Skip corrupted files
            continue
    return jsonify(projects_list)

@app.route('/api/projects/active', methods=['GET'])
def get_active_project():
    """Gets the active project's filename."""
    return jsonify({"active_filename": get_active_project_filename()})

@app.route('/api/projects/active', methods=['POST'])
def change_active_project():
    """Changes the active project to a different JSON file."""
    global project
    payload = request.json
    if not payload or "filename" not in payload:
        return jsonify({"error": "No filename specified"}), 400

    filename = payload["filename"]
    filepath = os.path.join(PROJECTS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "Project file not found"}), 404

    set_active_project_filename(filename)
    project = NovelProject(filepath)
    return jsonify({"status": "success", "active_filename": filename, "data": project.data})

@app.route('/api/projects/create', methods=['POST'])
def create_project():
    """Creates a new project file and sets it as active."""
    global project
    payload = request.json
    if not payload or "title" not in payload:
        return jsonify({"error": "Title is required"}), 400

    title = payload["title"].strip()
    if not title:
        return jsonify({"error": "Title cannot be empty"}), 400

    # Generate safe filename
    safe_title = "".join([c for c in title if c.isalnum() or c in " _-"]).rstrip()
    safe_title = safe_title.replace(" ", "_").lower()
    if not safe_title:
        safe_title = "unnamed_project"

    filename = f"{safe_title}.json"
    filepath = os.path.join(PROJECTS_DIR, filename)

    # If filename already exists, append a unique count
    counter = 1
    while os.path.exists(filepath):
        filename = f"{safe_title}_{counter}.json"
        filepath = os.path.join(PROJECTS_DIR, filename)
        counter += 1

    # Initialize clean project state
    new_proj = NovelProject(filepath)
    new_proj.data = new_proj.get_default_data()
    new_proj.data["settings"]["title"] = title
    new_proj.data["settings"]["overall_written"] = 0
    new_proj.data["settings"]["daily_written"] = 0
    new_proj.data["manuscript"] = []
    new_proj.data["plot"]["cards"] = []
    new_proj.data["characters"] = []
    new_proj.data["story_notes"] = []
    new_proj.save()

    set_active_project_filename(filename)
    project = new_proj

    return jsonify({"status": "success", "filename": filename, "data": project.data})

@app.route('/api/projects/delete', methods=['POST'])
def delete_project():
    """Deletes a project file securely."""
    global project
    payload = request.json or {}
    filename = payload.get("filename")
    if not filename:
        return jsonify({"error": "No filename specified"}), 400

    filepath = os.path.join(PROJECTS_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "Project file not found"}), 404

    try:
        os.remove(filepath)
    except Exception as e:
        return jsonify({"error": f"Failed to delete file: {str(e)}"}), 500

    # If we deleted the active project, find a new active project or create one
    active_fn = get_active_project_filename()
    if active_fn == filename or not os.path.exists(os.path.join(PROJECTS_DIR, active_fn)):
        if os.path.exists(ACTIVE_CONFIG_FILE):
            os.remove(ACTIVE_CONFIG_FILE)
        new_active = get_active_project_filename()
        project = NovelProject(os.path.join(PROJECTS_DIR, new_active))
        return jsonify({
            "status": "success",
            "switched": True,
            "active_filename": new_active,
            "data": project.data
        })

    return jsonify({
        "status": "success",
        "switched": False,
        "active_filename": active_fn
    })

@app.route('/api/locale/<lang>', methods=['GET'])
def get_locale(lang):
    """Loads and returns external translation files."""
    if lang not in ["en", "fr", "es", "ru"]:
        lang = "en"

    locale_path = os.path.join("locales", f"{lang}.json")
    try:
        with open(locale_path, 'r', encoding='utf-8') as f:
            translations = json.load(f)
        return jsonify(translations)
    except Exception as e:
        return jsonify({"error": f"Failed to load translations: {str(e)}"}), 500

@app.route('/api/synonyms', methods=['POST'])
def api_synonyms():
    """Returns a list of curated synonyms for a given word and language."""
    payload = request.get_json(silent=True) or {}
    word = payload.get("word", "").strip()
    lang = payload.get("lang", "fr").strip()

    if not word:
        return jsonify({"synonyms": []})

    syns = get_synonyms(word, lang)
    return jsonify({"synonyms": syns})

def clean_annotations(text):
    """Strips annotation span wrapping while preserving the inner annotated text."""
    import re
    return re.sub(r'<span[^>]*class="annotation-highlight[^>]*>(.*?)</span>', r'\1', text, flags=re.DOTALL)

def add_docx_formatted_paragraph(doc, text):
    """Helper to add paragraphs with basic <i>, <b> and <span style='font-variant: small-caps;'> styling."""
    paragraphs = text.split('\n')
    for para in paragraphs:
        p = doc.add_paragraph()
        parts = re.split(r'(<b><i>|</i></b>|<b>|</b>|<i>|</i>|<span style="font-variant: small-caps;">|</span>)', para)

        is_bold = False
        is_italic = False
        is_smallcaps = False

        for part in parts:
            if part == '<b><i>':
                is_bold = True
                is_italic = True
            elif part == '</i></b>':
                is_bold = False
                is_italic = False
            elif part == '<b>':
                is_bold = True
            elif part == '</b>':
                is_bold = False
            elif part == '<i>':
                is_italic = True
            elif part == '</i>':
                is_italic = False
            elif part == '<span style="font-variant: small-caps;">':
                is_smallcaps = True
            elif part == '</span>':
                is_smallcaps = False
            else:
                if part:
                    run = p.add_run(part)
                    run.bold = is_bold
                    run.italic = is_italic
                    if is_smallcaps:
                        run.font.small_caps = True

@app.route('/api/export', methods=['POST'])
def export_draft():
    """Compiles the entire manuscript and exports it in the chosen format."""
    global project
    try:
        payload = request.get_json(silent=True) or {}
        fmt = payload.get("format", "txt").lower().strip()

        title = project.data["settings"].get("title", "My Novel")
        safe_title = title.replace(' ', '_')

        # Compile clean manuscript text representation
        content_lines = [f"=== {title} ===\n\n"]
        for chap in project.data["manuscript"]:
            content_lines.append(f"\n--- {chap['title']} ---\n\n")
            for scene in chap.get("children", []):
                content_lines.append(f"[{scene['title']}]\n")
                content_lines.append(f"{clean_annotations(scene.get('content', ''))}\n\n")
        compiled_text = "".join(content_lines)

        if fmt == "txt":
            temp_file_path = "novel_export_temp.txt"
            with open(temp_file_path, 'w', encoding='utf-8') as f:
                f.write(compiled_text)
            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=f"{safe_title}_draft.txt",
                mimetype="text/plain"
            )

        elif fmt == "docx":
            import docx
            temp_file_path = "novel_export_temp.docx"
            doc = docx.Document()
            doc.add_heading(title, 0)

            for chap in project.data["manuscript"]:
                doc.add_heading(chap['title'], level=1)
                for scene in chap.get("children", []):
                    doc.add_heading(scene['title'], level=2)
                    add_docx_formatted_paragraph(doc, clean_annotations(scene.get('content', '')))

            doc.save(temp_file_path)
            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=f"{safe_title}_draft.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        elif fmt == "pdf":
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

            temp_file_path = "novel_export_temp.pdf"
            doc = SimpleDocTemplate(temp_file_path, pagesize=letter)
            story = []

            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Heading1'],
                alignment=TA_CENTER,
                fontSize=24,
                spaceAfter=20
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))

            heading1_style = ParagraphStyle(
                'Heading1Style',
                parent=styles['Heading2'],
                fontSize=18,
                spaceBefore=15,
                spaceAfter=10
            )
            heading2_style = ParagraphStyle(
                'Heading2Style',
                parent=styles['Heading3'],
                fontSize=14,
                spaceBefore=10,
                spaceAfter=6
            )
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['BodyText'],
                fontSize=11,
                leading=14,
                alignment=TA_JUSTIFY,
                spaceAfter=10
            )

            for chap in project.data["manuscript"]:
                story.append(Paragraph(chap['title'], heading1_style))
                for scene in chap.get("children", []):
                    story.append(Paragraph(scene['title'], heading2_style))
                    content = clean_annotations(scene.get('content', '')).replace('\n', '<br/>')
                    story.append(Paragraph(content, body_style))
                    story.append(Spacer(1, 10))

            doc.build(story)
            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=f"{safe_title}_draft.pdf",
                mimetype="application/pdf"
            )

        elif fmt == "odt":
            import zipfile
            import html
            temp_file_path = "novel_export_temp.odt"

            manifest_xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2">\n'
                ' <manifest:file-entry manifest:full-path="/" manifest:version="1.2" manifest:media-type="application/vnd.oasis.opendocument.text"/>\n'
                ' <manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>\n'
                ' <manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/>\n'
                ' <manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/>\n'
                '</manifest:manifest>\n'
            )

            meta_xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2">\n'
                ' <office:meta>\n'
                '  <meta:generator>Ecriture Novel Assistant</meta:generator>\n'
                f'  <meta:title>{html.escape(title)}</meta:title>\n'
                ' </office:meta>\n'
                '</office:document-meta>\n'
            )

            styles_xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2">\n'
                '</office:document-styles>\n'
            )

            body_xml_lines = []
            for chap in project.data["manuscript"]:
                body_xml_lines.append(f'<text:h text:outline-level="1">{html.escape(chap["title"])}</text:h>')
                for scene in chap.get("children", []):
                    body_xml_lines.append(f'<text:h text:outline-level="2">{html.escape(scene["title"])}</text:h>')
                    for para in clean_annotations(scene.get("content", "")).split("\n"):
                        if para.strip():
                            body_xml_lines.append(f'<text:p>{html.escape(para)}</text:p>')
            body_xml = "\n".join(body_xml_lines)

            content_xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" office:version="1.2">\n'
                ' <office:body>\n'
                '  <office:text>\n'
                f'   <text:h text:outline-level="1">{html.escape(title)}</text:h>\n'
                f'   {body_xml}\n'
                '  </office:text>\n'
                ' </office:body>\n'
                '</office:document-content>\n'
            )

            with zipfile.ZipFile(temp_file_path, 'w', zipfile.ZIP_DEFLATED) as o:
                o.writestr('mimetype', 'application/vnd.oasis.opendocument.text', compress_type=zipfile.ZIP_STORED)
                o.writestr('META-INF/manifest.xml', manifest_xml)
                o.writestr('meta.xml', meta_xml)
                o.writestr('styles.xml', styles_xml)
                o.writestr('content.xml', content_xml)

            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=f"{safe_title}_draft.odt",
                mimetype="application/vnd.oasis.opendocument.text"
            )

        elif fmt == "epub":
            import zipfile
            import html
            temp_file_path = "novel_export_temp.epub"

            container_xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n'
                '  <rootfiles>\n'
                '    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>\n'
                '  </rootfiles>\n'
                '</container>\n'
            )

            content_opf = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookID" version="2.0">\n'
                '  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">\n'
                f'    <dc:title>{html.escape(title)}</dc:title>\n'
                '    <dc:language>fr</dc:language>\n'
                '    <dc:creator>Ecriture Novel Assistant</dc:creator>\n'
                '    <dc:identifier id="BookID">urn:uuid:12345678-1234-1234-1234-123456789abc</dc:identifier>\n'
                '  </metadata>\n'
                '  <manifest>\n'
                '    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>\n'
                '    <item id="text" href="text.html" media-type="application/xhtml+xml"/>\n'
                '  </manifest>\n'
                '  <spine toc="ncx">\n'
                '    <itemref idref="text"/>\n'
                '  </spine>\n'
                '</package>\n'
            )

            toc_ncx = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<!DOCTYPE ncx PUBLIC "-//NISO//DTD NCX 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">\n'
                '<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">\n'
                '  <head>\n'
                '    <meta name="dtb:uid" content="urn:uuid:12345678-1234-1234-1234-123456789abc"/>\n'
                '    <meta name="dtb:depth" content="1"/>\n'
                '    <meta name="dtb:totalPageCount" content="0"/>\n'
                '    <meta name="dtb:maxPageNumber" content="0"/>\n'
                '  </head>\n'
                '  <docTitle>\n'
                f'    <text>{html.escape(title)}</text>\n'
                '  </docTitle>\n'
                '  <navMap>\n'
                '    <navPoint id="navPoint-1" playOrder="1">\n'
                '      <navLabel>\n'
                '        <text>Start</text>\n'
                '      </navLabel>\n'
                '      <content src="text.html"/>\n'
                '    </navPoint>\n'
                '  </navMap>\n'
                '</ncx>\n'
            )

            html_body_lines = []
            for chap in project.data["manuscript"]:
                html_body_lines.append(f'<h2>{html.escape(chap["title"])}</h2>')
                for scene in chap.get("children", []):
                    html_body_lines.append(f'<h3>{html.escape(scene["title"])}</h3>')
                    for para in clean_annotations(scene.get("content", "")).split("\n"):
                        if para.strip():
                            html_body_lines.append(f'<p>{html.escape(para)}</p>')
            html_body = "\n".join(html_body_lines)

            text_html = (
                '<?xml version="1.0" encoding="utf-8"?>\n'
                '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">\n'
                '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                '<head>\n'
                f'  <title>{html.escape(title)}</title>\n'
                '</head>\n'
                '<body>\n'
                f'  <h1>{html.escape(title)}</h1>\n'
                f'  {html_body}\n'
                '</body>\n'
                '</html>\n'
            )

            with zipfile.ZipFile(temp_file_path, 'w', zipfile.ZIP_DEFLATED) as o:
                o.writestr('mimetype', 'application/epub+zip', compress_type=zipfile.ZIP_STORED)
                o.writestr('META-INF/container.xml', container_xml)
                o.writestr('OEBPS/content.opf', content_opf)
                o.writestr('OEBPS/toc.ncx', toc_ncx)
                o.writestr('OEBPS/text.html', text_html)

            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=f"{safe_title}_draft.epub",
                mimetype="application/epub+zip"
            )

        elif fmt == "mobi":
            import struct
            temp_file_path = "novel_export_temp.mobi"

            title_bytes = title.encode('utf-8')[:31].ljust(32, b'\0')
            num_records = 3

            pdb_header = struct.pack(
                '>32sHHIIIIII4s4sIIH',
                title_bytes,
                0, # attributes
                0, # version
                0, 0, 0, 0, # dates
                0, # app info
                0, # sort info
                b'BOOK',
                b'MOBI',
                0, # unique id seed
                0, # next record
                num_records
            )

            text_bytes = compiled_text.encode('utf-8')
            text_len = len(text_bytes)

            # Rec0 is the PalmDOC/Mobi description header:
            # - compression: 1 (none)
            # - unused: 0
            # - text length
            # - record count: 1 (text records count)
            # - record size: 4096
            # - encryption: 0
            rec0 = struct.pack('>HHIIHH', 1, 0, text_len, num_records - 1, 4096, 0)
            rec1 = text_bytes
            rec2 = b'\xe9\x8e\r\n' # EOF

            offset0 = 78 + (num_records * 8) + 2
            offset1 = offset0 + len(rec0)
            offset2 = offset1 + len(rec1)

            rec_info = b''
            rec_info += struct.pack('>II', offset0, 0)
            rec_info += struct.pack('>II', offset1, 2)
            rec_info += struct.pack('>II', offset2, 4)

            mobi_bytes = pdb_header + rec_info + b'\0\0' + rec0 + rec1 + rec2

            with open(temp_file_path, 'wb') as f:
                f.write(mobi_bytes)

            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=f"{safe_title}_draft.mobi",
                mimetype="application/x-mobipocket-ebook"
            )

        else:
            return jsonify({"error": f"Unsupported format: {fmt}"}), 400

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to export: {str(e)}"}), 500

@app.route('/api/ai', methods=['POST'])
def handle_ai_tool():
    """Handles contextual AI writing tools: Describe, Rewrite, and Expand."""
    from ai_prompts import DESCRIBE_PROMPT, REWRITE_PROMPT, EXPAND_PROMPT, SHOW_DONT_TELL_PROMPT, SENSORY_PROMPT, POV_PROMPT, COMPLICATIONS_PROMPT, NAMES_PROMPT

    payload = request.json or {}
    tool = payload.get("tool", "describe").lower().strip()
    style = payload.get("style", "elegant").lower().strip()
    text = payload.get("text", "").strip()
    inject_lore = payload.get("inject_lore_context", True)
    scene_id = payload.get("scene_id")
    lang = payload.get("lang", "fr")

    if not text and tool not in ["names", "complications"]:
        return jsonify({"error": "No text selected"}), 400

    # Match the appropriate system prompt
    if tool == "describe":
        system_prompt = DESCRIBE_PROMPT
    elif tool == "rewrite":
        system_prompt = REWRITE_PROMPT.format(style=style)
    elif tool == "pov":
        system_prompt = POV_PROMPT.format(style=style)
    elif tool == "expand":
        system_prompt = EXPAND_PROMPT
    elif tool == "show_dont_tell":
        system_prompt = SHOW_DONT_TELL_PROMPT
    elif tool == "sensory":
        system_prompt = SENSORY_PROMPT
    elif tool == "complications":
        system_prompt = COMPLICATIONS_PROMPT
    elif tool == "names":
        system_prompt = NAMES_PROMPT.format(style=style)
    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

    mapping = {'fr': 'French', 'es': 'Spanish', 'ru': 'Russian'}
    lang_name = mapping.get(lang, 'English')
    lang_instruction = f"Respond strictly in this language: {lang_name}."
    system_prompt += lang_instruction

    # Prepend Lore context to the system prompt if enabled and context is found
    if inject_lore and scene_id:
        lore_ctx = get_scene_context(scene_id)
        if lore_ctx:
            system_prompt = (
                f"{lore_ctx}\n\n"
                "Consignes de l'assistant :\n"
                f"{system_prompt}"
            )

    # Build chat messages payload
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": text}
    ]

    selected_model = payload.get("model", "llama3").strip()
    temperature = payload.get("temperature", 0.7)

    # Call Gemma via unified Client
    try:
        res = ai_client.generate_chat(messages, model=selected_model, temperature=temperature, timeout=15)
        return jsonify({
            "status": "success",
            "message": res["message"],
            "model": res["model"]
        })
    except Exception:
        # Fallback using unified client fallbacks
        lang = "fr" if any(word in text.lower() for word in ["le", "la", "les", "une", "un", "est", "et", "de", "je", "tu", "il"]) else "en"
        simulated_output = ai_client.get_fallback_response(tool, text, style=style, lang=lang)
        return jsonify({
            "status": "offline_fallback",
            "message": simulated_output,
            "model": "Simulation"
        })

@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    """Proxy or fallback endpoint for a local Gemma AI assistant."""
    payload = request.json or {}
    messages = payload.get("messages", [])
    selected_model = payload.get("model", "llama3").strip()
    temperature = payload.get("temperature", 0.7)
    inject_lore = payload.get("inject_lore_context", True)
    scene_id = payload.get("scene_id")
    lang = payload.get("lang", "fr")

    mapping = {'fr': 'French', 'es': 'Spanish', 'ru': 'Russian'}
    lang_name = mapping.get(lang, 'English')
    lang_instruction = f"Respond strictly in this language: {lang_name}."
    system_msgs = [lang_instruction]

    # Prepend system message with lore context if enabled and context is found
    if inject_lore and scene_id:
        lore_ctx = get_scene_context(scene_id)
        if lore_ctx:
            if lang == 'fr':
                system_msgs.append(f"Voici des informations sur le contexte et le Lore de la scène en cours. Intègre et respecte ces éléments si nécessaire dans vos réponses :\n\n{lore_ctx}")
            elif lang == 'es':
                system_msgs.append(f"Aquí hay información sobre el contexto y la tradición de la escena actual. Integra y respeta estos elementos si es necesario en tus respuestas:\n\n{lore_ctx}")
            elif lang == 'ru':
                system_msgs.append(f"Здесь представлена информация о контексте и лоре текущей сцены. Интегрируйте и учитывайте эти элементы при необходимости в своих ответах:\n\n{lore_ctx}")
            else:
                system_msgs.append(f"Here is information on the context and lore of the current scene. Integrate and respect these elements if necessary in your answers:\n\n{lore_ctx}")

    messages.insert(0, {"role": "system", "content": "\n\n".join(system_msgs)})

    # Call Gemma via unified client
    try:
        res = ai_client.generate_chat(messages, model=selected_model, temperature=temperature, timeout=15)
        return jsonify({
            "status": "success",
            "message": res["message"],
            "model": res["model"]
        })
    except Exception:
        # Graceful fallback simulation
        simulation = ai_client.get_fallback_response("chat", messages)
        return jsonify({
            "status": "offline_fallback",
            "message": simulation,
            "model": "Simulation"
        })

@app.route('/api/ai/status', methods=['GET'])
def get_ai_status():
    """Checks if Gemma is installed and running locally."""
    ai_status = ai_client.check_status()
    models = ai_client.get_models()

    import shutil
    import psutil

    total, used, free = shutil.disk_usage("/")
    free_gb = free / (1024**3)

    mem = psutil.virtual_memory()
    ram_gb = mem.total / (1024**3)

    import sys
    os_name = sys.platform

    sys_info = {
        "free_disk_gb": round(free_gb, 2),
        "total_ram_gb": round(ram_gb, 2),
        "os": os_name
    }

    if ai_status.get("status") == "online":
        return jsonify({
            "status": "online",
            "installed": True,
            "models": models,
            "sys_info": sys_info
        })
    else:
        return jsonify({
            "status": "offline",
            "installed": False,
            "error": ai_status.get("error", "Unknown error"),
            "traceback": ai_status.get("traceback", ""),
            "sys_info": sys_info
        })


CURRENT_VERSION = "1.0.0"
GITHUB_REPO = "neomars/ecriture"
API_URL = f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest"

def get_os_keyword():
    system = platform.system().lower()
    if system == "windows":
        return "windows"
    elif system == "darwin":
        return "macos"
    return "linux"

@app.route('/api/check_updates', methods=['GET'])
def check_updates():
    try:
        req = urllib.request.Request(
            API_URL,
            headers={"User-Agent": "Ecriture-App-Updater"}
        )
        with urllib.request.urlopen(req, timeout=5) as response:
            data = json.loads(response.read().decode("utf-8"))

        latest_tag = data.get("tag_name", "").lstrip("v")

        if version.parse(latest_tag) > version.parse(CURRENT_VERSION):
            os_keyword = get_os_keyword()
            download_url = data.get("html_url")

            for asset in data.get("assets", []):
                asset_name = asset.get("name", "").lower()
                if os_keyword in asset_name:
                    download_url = asset.get("browser_download_url")
                    break

            return jsonify({
                "update_available": True,
                "current_version": CURRENT_VERSION,
                "latest_version": latest_tag,
                "download_url": download_url,
                "release_notes": data.get("body", ""),
                "release_page": data.get("html_url")
            })
    except Exception as e:
        print(f"[Updater] Impossible de vérifier les mises à jour: {e}")

    return jsonify({
        "update_available": False,
        "current_version": CURRENT_VERSION
    })


INSTALL_STATE = {
    "status": "idle",
    "message": "",
    "progress": 0
}


def _install_gemma_thread(lang="fr"):
    global INSTALL_STATE
    import time
    from tqdm.auto import tqdm

    def get_str(key):
        import json
        import os
        # Extremely basic fallback translation
        filepath = f"locales/{lang}.json"
        if not os.path.exists(filepath):
            filepath = "locales/fr.json"

        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f).get(key, key)

    class CustomTqdm(tqdm):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._last_update_time = time.time()

        def update(self, n=1):
            super().update(n)
            current_time = time.time()
            if current_time - self._last_update_time > 1.0 or self.n == self.total:
                self._last_update_time = current_time
                if self.total and self.total > 0:
                    percent = (self.n / self.total) * 100
                    INSTALL_STATE["progress"] = min(99, int(percent))

                    elapsed = current_time - self.start_t
                    if elapsed > 0 and self.n > 0:
                        speed = self.n / elapsed
                        remaining_bytes = self.total - self.n
                        eta_seconds = remaining_bytes / speed
                        INSTALL_STATE["eta"] = int(eta_seconds)
                    else:
                        INSTALL_STATE["eta"] = None

    INSTALL_STATE["status"] = "installing_model"
    INSTALL_STATE["message"] = get_str("gemma_installing_model")
    INSTALL_STATE["progress"] = 0

    try:
        import urllib.request
        import urllib.error
        import os
        import time

        url = "https://neomars.freeboxos.fr:3535/share/SXyfQB1empSwlfSU/gemma-2-2b-it-Q8_0.gguf"
        model_dir = os.path.join(os.path.expanduser("~"), ".cache", "ecriture")
        os.makedirs(model_dir, exist_ok=True)
        model_path = os.path.join(model_dir, "gemma-2-2b-it-Q8_0.gguf")

        # Set up a context to bypass potential ssl errors if required, but default to normal.
        import ssl
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, context=ctx) as response, open(model_path, 'wb') as out_file:
            total_size = int(response.info().get('Content-Length').strip())

            t = CustomTqdm(total=total_size, unit='iB', unit_scale=True)
            t.start_t = time.time()

            chunk_size = 1024 * 1024 # 1MB chunks
            while True:
                chunk = response.read(chunk_size)
                if not chunk:
                    break
                out_file.write(chunk)
                t.update(len(chunk))

            t.close()

        INSTALL_STATE["status"] = "done"
        INSTALL_STATE["message"] = get_str("gemma_install_success")
        INSTALL_STATE["progress"] = 100

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error installing model: {e}")
        INSTALL_STATE["status"] = "error"
        err_msg = get_str("gemma_install_error").replace("{error}", str(e))
        INSTALL_STATE["message"] = err_msg


@app.route('/api/ai/install_engine', methods=['POST'])
def install_engine():
    try:
        global INSTALL_STATE
        if INSTALL_STATE["status"] in ["installing_model"]:
            return jsonify({"status": "success", "message": "Installation déjà en cours"})

        import threading
        lang = "fr" # could parse from headers/cookies if needed, default to fr.
        thread = threading.Thread(target=_install_gemma_thread, args=(lang,))
        thread.start()

        return jsonify({"status": "success", "message": "Installation started"})
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500

@app.route('/api/ai/install_status', methods=['GET'])
def get_install_status():
    global INSTALL_STATE
    return jsonify(INSTALL_STATE)


@app.route('/api/ai/models', methods=['GET'])
def get_ai_models():
    """Queries the local Gemma API to fetch installed models."""
    models = ai_client.get_models()
    if models:
        return jsonify({
            "status": "success",
            "models": models
        })
    else:
        return jsonify({
            "status": "offline",
            "models": []
        })

@app.route('/api/relecture/ai', methods=['POST'])
def api_relecture_ai():
    """AI relecture assistance: Style & Prose or Cohérence."""
    payload = request.json or {}
    category = payload.get("category", "style").lower().strip()
    text = payload.get("text", "").strip()
    selected_model = payload.get("model", "llama3").strip()
    temperature = payload.get("temperature", 0.7)
    lang = payload.get("lang", "fr").strip()

    if not text:
        empty_msg = {"fr": "Texte vide", "es": "Texto vacío", "ru": "Текст пуст"}.get(lang, "Empty text")
        return jsonify({"feedback": empty_msg, "status": "empty"})

    # Formulate prompt based on category and language
    if category == "style":
        if lang == "fr":
            system_prompt = (
                "Tu es un relecteur professionnel de romans et correcteur littéraire de style et prose.\n"
                "Analyse le texte suivant et donne des retours constructifs détaillés.\n"
                "Suggère des améliorations précises de vocabulaire, de rythme des phrases, de style, "
                "de fluidité et des reformulations d'échantillons de texte s'il y a lieu.\n"
                "Réponds en français."
            )
        elif lang == "es":
            system_prompt = (
                "Eres un corrector profesional de novelas y editor literario de estilo y prosa.\n"
                "Analiza el siguiente texto y proporciona comentarios constructivos detallados.\n"
                "Sugiere mejoras precisas de vocabulario, ritmo de oraciones, estilo, "
                "fluidez y reescrituras de muestra donde corresponda.\n"
                "Responde en español."
            )
        elif lang == "ru":
            system_prompt = (
                "Вы профессиональный корректор романов и литературный редактор стиля и прозы.\n"
                "Проанализируйте следующий текст и дайте подробные конструктивные отзывы.\n"
                "Предложите точные улучшения словарного запаса, ритма предложений, стиля, "
                "текучести и образцы перефразирования текста, где это применимо.\n"
                "Отвечайте на русском языке."
            )
        else:
            system_prompt = (
                "You are a professional novel proofreader and copyeditor of style and prose.\n"
                "Analyze the following text and provide detailed constructive feedback.\n"
                "Suggest precise improvements for vocabulary, sentence pacing, style, "
                "flow, and sample rewrites where applicable.\n"
                "Respond in English."
            )
    elif category == "worldbuilding":
        lore_context = payload.get("lore_context", "Aucun contexte fourni.")
        if lang == "fr":
            from ai_prompts import LORE_COHERENCE_PROMPT_FR
            system_prompt = LORE_COHERENCE_PROMPT_FR.format(lore_context=lore_context)
        elif lang == "es":
            from ai_prompts import LORE_COHERENCE_PROMPT_ES
            system_prompt = LORE_COHERENCE_PROMPT_ES.format(lore_context=lore_context)
        elif lang == "ru":
            from ai_prompts import LORE_COHERENCE_PROMPT_RU
            system_prompt = LORE_COHERENCE_PROMPT_RU.format(lore_context=lore_context)
        else:
            from ai_prompts import LORE_COHERENCE_PROMPT_EN
            system_prompt = LORE_COHERENCE_PROMPT_EN.format(lore_context=lore_context)
    else: # coherence
        if lang == "fr":
            system_prompt = (
                "Tu es un relecteur professionnel de romans et conseiller en cohérence narrative.\n"
                "Analyse le texte suivant pour en évaluer la cohérence logique, les motivations et actions des personnages, "
                "la pertinence temporelle et spatiale, et signale toute anomalie ou incohérence flagrante.\n"
                "Réponds en français."
            )
        else:
            system_prompt = (
                "You are a professional novel proofreader and narrative coherence consultant.\n"
                "Analyze the following text to evaluate logical consistency, character motivations and actions, "
                "temporal and spatial sense, and flag any logical fallacies or glaring inconsistencies.\n"
                "Respond in English."
            )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": text}
    ]

    # Call Gemma via unified client
    try:
        res = ai_client.generate_chat(messages, model=selected_model, temperature=temperature, timeout=25)
        return jsonify({
            "status": "success",
            "feedback": res["message"],
            "model": res["model"]
        })
    except Exception:
        # Fallback using unified client fallbacks
        simulated_feedback = ai_client.get_fallback_response(f"relecture_{category}", text, lang=lang)
        return jsonify({
            "status": "offline_fallback",
            "feedback": simulated_feedback,
            "model": "Simulation"
        })

@app.route('/api/backups/choose_directory', methods=['POST'])
def choose_directory():
    """Opens a native OS folder picker dialog to select a backup directory."""
    import tkinter as tk
    from tkinter import filedialog
    import os

    try:
        root = tk.Tk()
        root.withdraw()
        # Bring the dialog window to the front
        root.wm_attributes('-topmost', 1)
        folder = filedialog.askdirectory(title="Sélectionner le répertoire de destination des sauvegardes")
        root.destroy()
        if folder:
            return jsonify({"status": "success", "folder_path": os.path.abspath(folder)})
        else:
            return jsonify({"status": "cancelled", "folder_path": ""})
    except Exception as e:
        # Graceful fallback for headless environments
        return jsonify({
            "status": "headless_fallback",
            "error": str(e),
            "folder_path": os.path.abspath("./backups")
        })

@app.route('/api/backups/local/create', methods=['POST'])
def local_backup_create():
    payload = request.json or {}
    folder_path = payload.get("folder_path", "").strip()
    frequency = payload.get("frequency", "manual").strip() # "daily", "weekly", "monthly", "manual"
    if not folder_path:
        return jsonify({"error": "No folder path provided"}), 400

    import os
    import json
    from datetime import datetime

    # Create directory if not exists
    try:
        os.makedirs(folder_path, exist_ok=True)
    except Exception as e:
        return jsonify({"error": f"Failed to create directory: {str(e)}"}), 500

    # Retrieve current active project data
    current_data = project.data if project else None
    if not current_data:
        return jsonify({"error": "No active project data to back up"}), 400

    # Generate filename
    clean_title = "".join(c for c in current_data.get("settings", {}).get("title", "roman") if c.isalnum() or c in (' ', '_', '-')).rstrip()
    clean_title = clean_title.replace(' ', '_')
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"backup_{clean_title}_{timestamp}_{frequency}.json"
    full_path = os.path.join(folder_path, filename)

    try:
        with open(full_path, "w", encoding="utf-8") as f:
            json.dump(current_data, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "filename": filename, "full_path": full_path})
    except Exception as e:
        return jsonify({"error": f"Failed to write backup file: {str(e)}"}), 500

@app.route('/api/quit', methods=['POST'])
def quit_app():
    """Shuts down the backend web server gracefully."""
    import os
    import signal
    os.kill(os.getpid(), signal.SIGINT)
    return jsonify({"status": "shutdown"})


@app.route('/api/backups/local/list', methods=['GET'])
def local_backup_list():
    """List all available backups for the current project."""
    if not project.data or not project.data.get('id'):
        return jsonify({"error": "No active project"}), 400

    # Read path from query params, fallback to project settings, fallback to default
    backup_dir = request.args.get('path')
    if not backup_dir:
        backup_dir = project.data.get("settings", {}).get("backup_config", {}).get("folder_path", os.path.abspath("./backups"))

    if not os.path.exists(backup_dir):
        return jsonify({"backups": []})

    pid = project.data['id']
    backups = []

    for f in os.listdir(backup_dir):
        if f.endswith('.json') and f.startswith('backup_'):
            filepath = os.path.join(backup_dir, f)
            try:
                import json
                with open(filepath, 'r', encoding='utf-8') as f_json:
                    data = json.load(f_json)
                    if data.get('id') == pid:
                        backups.append({
                            "filename": f,
                            "path": filepath,
                            "timestamp": os.path.getmtime(filepath),
                            "size": os.path.getsize(filepath)
                        })
            except Exception:
                continue

    backups.sort(key=lambda x: x['timestamp'], reverse=True)
    return jsonify({"backups": backups})

@app.route('/api/backups/local/restore', methods=['POST'])
def local_backup_restore():
    """Restore a specific backup file."""
    data = request.json
    filename = data.get("filename")
    path = data.get("path")
    if not filename:
        return jsonify({"error": "No filename provided"}), 400

    backup_dir = path if path else project.data.get("settings", {}).get("backup_config", {}).get("folder_path", os.path.abspath("./backups"))
    # Sanitize filename to prevent path traversal
    from werkzeug.utils import secure_filename
    filename = secure_filename(filename)
    filepath = os.path.join(backup_dir, filename)

    if not os.path.exists(filepath):
        return jsonify({"error": "Backup file not found"}), 404

    try:
        import json
        with open(filepath, 'r', encoding='utf-8') as f:
            backup_data = json.load(f)

        project.data = backup_data

        if project.filepath:
            with open(project.filepath, 'w', encoding='utf-8') as f:
                json.dump(project.data, f, indent=4, ensure_ascii=False)

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/import/document', methods=['POST'])
def import_document():
    """Import a document and overwrite the current project's manuscript."""
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    from werkzeug.utils import secure_filename
    import zipfile
    from bs4 import BeautifulSoup
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

    if ext not in ['txt', 'docx', 'odt', 'epub']:
        return jsonify({"error": "Unsupported file format"}), 400

    temp_dir = os.path.abspath('./temp_imports')
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, filename)
    file.save(temp_path)

    text_content = ""
    try:
        if ext == 'txt':
            with open(temp_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
        elif ext == 'docx':
            from docx import Document as DocxDocument
            doc = DocxDocument(temp_path)
            text_content = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == 'odt':
            with zipfile.ZipFile(temp_path, 'r') as z:
                xml_content = z.read('content.xml')
                soup = BeautifulSoup(xml_content, 'xml')
                paragraphs = soup.find_all('text:p')
                text_content = "\n\n".join([p.text for p in paragraphs if p.text.strip()])
        elif ext == 'epub':
            with zipfile.ZipFile(temp_path, 'r') as z:
                for item in z.namelist():
                    if item.endswith('.html') or item.endswith('.xhtml'):
                        html_content = z.read(item)
                        soup = BeautifulSoup(html_content, 'html.parser')
                        text_content += "\n\n" + soup.get_text(separator='\n')

        text_content = text_content.strip()
        html_content = text_content.replace('\n', '<br>')

        project.data["manuscript"] = [{
            "id": "chap-imported",
            "title": "Chapitre Importé",
            "type": "corps",
            "children": [{
                "id": "scene-imported",
                "title": "Scène Importée",
                "content": html_content
            }]
        }]

        if project.filepath:
            import json
            with open(project.filepath, 'w', encoding='utf-8') as f:
                json.dump(project.data, f, indent=4, ensure_ascii=False)

    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({"error": str(e)}), 500

    if os.path.exists(temp_path):
        os.remove(temp_path)

    return jsonify({"success": True, "message": "Import successful"})



@app.route('/api/ai/pull', methods=['GET'])
def pull_ai_model():
    """Pulls an Gemma model and streams progress via Server-Sent Events (SSE)."""
    model_name = request.args.get('model', 'gemma4:latest')

    def generate():
        import json
        import urllib.request
        try:
            req = urllib.request.Request(
                "http://localhost:11434/api/pull",
                data=json.dumps({"name": model_name, "stream": True}).encode('utf-8'),
                headers={'Content-Type': 'application/json'},
                method="POST"
            )
            with urllib.request.urlopen(req) as response:
                for line in response:
                    if line:
                        decoded_line = line.decode('utf-8')
                        try:
                            data = json.loads(decoded_line)
                            yield f"data: {json.dumps(data)}\n\n"
                        except json.JSONDecodeError:
                            continue
                yield f"data: {json.dumps({'status': 'success'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    from flask import Response
    return Response(generate(), mimetype='text/event-stream')



@app.route('/api/ai/extract_characters', methods=['POST'])
def api_extract_characters():
    """Extract character information from text using AI."""
    payload = request.json or {}
    text = payload.get("text", "").strip()
    selected_model = payload.get("model", "llama3").strip()
    temperature = payload.get("temperature", 0.1)  # Low temp for data extraction
    lang = payload.get("lang", "fr")

    if not text:
        return jsonify({"status": "empty", "characters": []})

    from ai_prompts import EXTRACT_LORE_PROMPT

    mapping = {'fr': 'French', 'es': 'Spanish', 'ru': 'Russian'}
    lang_name = mapping.get(lang, 'English')
    lang_instruction = f"Respond strictly in this language: {lang_name}."
    messages = [
        {"role": "system", "content": EXTRACT_LORE_PROMPT + lang_instruction},
        {"role": "user", "content": text}
    ]

    try:
        res = ai_client.generate_chat(messages, model=selected_model, temperature=temperature, timeout=30)
        # Try to parse the JSON
        import json
        import re

        content = res["message"]
        # Find JSON array using regex in case model wrapped it in backticks
        match = re.search(r'\[.*\]', content, re.DOTALL)
        if match:
            json_str = match.group(0)
            try:
                characters = json.loads(json_str)
                return jsonify({"status": "success", "characters": characters})
            except json.JSONDecodeError:
                pass

        return jsonify({"status": "error", "message": "Failed to parse JSON response"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})


if __name__ == "__main__":
    import webbrowser
    import threading
    import os

    def open_browser():
        webbrowser.open("http://127.0.0.1:5000")

    # Only launch browser once when in debug reloader
    if not app.debug or os.environ.get("WERKZEUG_RUN_MAIN") == "true":
        threading.Timer(1.5, open_browser).start()

    # Start the local development web server on port 5000
    app.run(host="0.0.0.0", port=5000, debug=True)
